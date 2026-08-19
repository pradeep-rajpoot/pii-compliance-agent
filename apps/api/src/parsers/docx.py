from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from src.models.block import TextBlock
from src.models.enums import FileType
from src.models.locator import DocxLocator
from src.parsers.common import ParsedDocument


def compute_run_offsets(paragraph: Paragraph) -> list[tuple[int, int, int]]:
    """Return [(run_index, start_char, end_char), ...] mapping each run's
    text to its character range within `paragraph.text` (which is simply the
    concatenation of all run texts). Used both to build detection-time meta
    and, freshly recomputed, by the correction agent to splice masked text
    back into the correct run(s) -- including spans that straddle a run
    boundary."""

    offsets: list[tuple[int, int, int]] = []
    cursor = 0
    for run_index, run in enumerate(paragraph.runs):
        length = len(run.text)
        offsets.append((run_index, cursor, cursor + length))
        cursor += length
    return offsets


def parse(path: Path) -> ParsedDocument:
    document = Document(str(path))

    blocks: list[TextBlock] = []
    run_offsets: dict[str, list[tuple[int, int, int]]] = {}

    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text
        if not text.strip():
            continue
        block_id = f"p{i}"
        blocks.append(
            TextBlock(
                id=block_id,
                text=text,
                locator=DocxLocator(paragraph=i),
            )
        )
        run_offsets[block_id] = compute_run_offsets(paragraph)

    return ParsedDocument(
        file_type=FileType.DOCX,
        blocks=blocks,
        meta={"run_offsets": run_offsets},
    )
