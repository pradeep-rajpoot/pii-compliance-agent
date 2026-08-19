import time

from docx import Document

from src.parsers import docx as docx_parser


def _poll_until(client, job_id, terminal_statuses, timeout_s=5.0):
    deadline = time.monotonic() + timeout_s
    last = None
    while time.monotonic() < deadline:
        resp = client.get(f"/api/jobs/{job_id}")
        assert resp.status_code == 200
        last = resp.json()
        if last["status"] in terminal_statuses:
            return last
        time.sleep(0.02)
    raise AssertionError(f"job never reached {terminal_statuses}, last={last}")


def test_full_detect_correct_download_flow(client, mock_bedrock, fixtures_dir):
    parsed = docx_parser.parse(fixtures_dir / "sample.docx")
    p0 = next(b for b in parsed.blocks if b.id == "p0")
    email = "jane.doe@example.com"
    start = p0.text.index(email)
    end = start + len(email)

    mock_bedrock.set_responses(
        [
            [
                {
                    "block_id": "p0",
                    "matched_text": email,
                    "pii_type": "email",
                    "start_offset": start,
                    "end_offset": end,
                    "confidence": 0.97,
                }
            ]
        ]
    )

    with open(fixtures_dir / "sample.docx", "rb") as f:
        resp = client.post(
            "/api/jobs/detect",
            files={
                "file": (
                    "sample.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert resp.status_code == 202
    body = resp.json()
    assert body["status"] == "queued"
    job_id = body["job_id"]

    detected = _poll_until(client, job_id, {"detected", "failed"})
    assert detected["status"] == "detected", detected
    assert detected["file_type"] == "docx"
    assert len(detected["detections"]) == 1
    d = detected["detections"][0]
    assert d["pii_type"] == "email"
    assert d["value"] == email

    resp = client.post(f"/api/jobs/{job_id}/correct")
    assert resp.status_code == 202
    assert resp.json()["status"] == "correcting"

    corrected = _poll_until(client, job_id, {"corrected", "failed"})
    assert corrected["status"] == "corrected", corrected

    resp = client.get(f"/api/jobs/{job_id}/download")
    assert resp.status_code == 200
    assert "attachment" in resp.headers["content-disposition"]
    assert "sample-pii-safe.docx" in resp.headers["content-disposition"]

    import io

    doc = Document(io.BytesIO(resp.content))
    assert email not in doc.paragraphs[0].text
    assert "X" * len(email) in doc.paragraphs[0].text
    assert doc.paragraphs[1].text == "Plain paragraph with no PII, just filler text."
