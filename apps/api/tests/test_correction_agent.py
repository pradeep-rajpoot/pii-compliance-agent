import shutil
from pathlib import Path

import openpyxl
import pdfplumber
from docx import Document

from src.agents.correction_agent import (
    _correct_csv,
    _correct_docx,
    _correct_json,
    _correct_pdf,
    _correct_xlsx,
)
from src.jobs.store import JobRecord
from src.models.detection import Detection
from src.models.enums import FileType, JobStatus
from src.parsers import csv as csv_parser
from src.parsers import docx as docx_parser
from src.parsers import json as json_parser
from src.parsers import pdf as pdf_parser


def _make_record(tmp_path: Path, fixtures_dir: Path, src_name: str, file_type: FileType) -> JobRecord:
    original_path = tmp_path / src_name
    shutil.copy(fixtures_dir / src_name, original_path)
    return JobRecord(
        job_id="job1",
        status=JobStatus.CORRECTING,
        file_type=file_type,
        original_path=str(original_path),
        original_filename=src_name,
        temp_dir=str(tmp_path),
    )


def _det(block_id: str, pii_type: str, value: str, start: int, confidence: float = 0.9) -> Detection:
    return Detection(
        id="d1",
        block_id=block_id,
        pii_type=pii_type,
        value=value,
        start_offset=start,
        end_offset=start + len(value),
        confidence=confidence,
    )


def test_correct_xlsx_masks_target_cell_and_preserves_formula(tmp_path, fixtures_dir):
    record = _make_record(tmp_path, fixtures_dir, "sample.xlsx", FileType.XLSX)
    email = "john.smith@example.com"
    det = _det("Sheet1!B2", "email", email, 0)

    out_path = _correct_xlsx(record, {"Sheet1!B2": [det]})

    wb = openpyxl.load_workbook(str(out_path))
    ws = wb["Sheet1"]
    assert ws["B2"].value == "X" * len(email)
    assert ws["A2"].value == "John Smith"  # untouched
    assert str(ws["B3"].value).startswith("=")  # formula preserved


def test_correct_csv_masks_target_cell_and_preserves_dialect(tmp_path, fixtures_dir):
    parsed = csv_parser.parse(fixtures_dir / "sample.csv")
    record = _make_record(tmp_path, fixtures_dir, "sample.csv", FileType.CSV)
    record.parse_meta = {"dialect": parsed.meta["dialect"]}

    email = "alice.example@example.com"
    det = _det("r1c1", "email", email, 0)

    out_path = _correct_csv(record, {"r1c1": [det]})

    raw_bytes = out_path.read_bytes()
    text = raw_bytes.decode("utf-8")
    assert email not in text
    assert "X" * len(email) in text
    assert "Alice Example" in text  # untouched neighboring cell
    assert "Bob Sample" in text  # untouched other row
    assert b"\r\n" in raw_bytes  # dialect's lineterminator preserved


def test_correct_docx_masks_span_that_straddles_run_boundary(tmp_path, fixtures_dir):
    parsed = docx_parser.parse(fixtures_dir / "sample.docx")
    p0 = next(b for b in parsed.blocks if b.id == "p0")
    email = "jane.doe@example.com"
    start = p0.text.index(email)

    record = _make_record(tmp_path, fixtures_dir, "sample.docx", FileType.DOCX)
    det = _det("p0", "email", email, start)

    out_path = _correct_docx(record, {"p0": [det]})

    doc = Document(str(out_path))
    para0 = doc.paragraphs[0]
    assert email not in para0.text
    assert "X" * len(email) in para0.text
    assert para0.text.startswith("Contact: ")
    assert para0.text.endswith("please call soon.")

    # Formatting outside the masked span must survive: the paragraph's
    # first run was originally bold, covering chars [0, 20). The prefix
    # "Contact: " sits entirely inside that bold run.
    assert doc.paragraphs[0].runs[0].bold is True

    # The untouched second paragraph must be completely unaffected.
    assert doc.paragraphs[1].text == "Plain paragraph with no PII, just filler text."


def test_correct_pdf_masks_text_and_rerenders(tmp_path, fixtures_dir):
    parsed = pdf_parser.parse(fixtures_dir / "sample.pdf")
    page1 = next(b for b in parsed.blocks if b.id == "page-1")
    email = "jane.doe@example.com"
    phone = "555-010-1234"
    ssn = "000-00-0000"

    dets = [
        _det("page-1", "email", email, page1.text.index(email)),
        _det("page-1", "phone", phone, page1.text.index(phone)),
        _det("page-1", "ssn", ssn, page1.text.index(ssn)),
    ]
    # give each a distinct id since _det always uses "d1"
    for i, d in enumerate(dets, start=1):
        d.id = f"d{i}"

    record = _make_record(tmp_path, fixtures_dir, "sample.pdf", FileType.PDF)

    out_path = _correct_pdf(record, {"page-1": dets})

    with pdfplumber.open(str(out_path)) as pdf:
        full_text = "\n".join(p.extract_text() or "" for p in pdf.pages)

    assert email not in full_text
    assert phone not in full_text
    assert ssn not in full_text
    assert "widgets" in full_text  # page 2 filler content untouched


def test_correct_json_masks_nested_and_array_leaves(tmp_path, fixtures_dir):
    parsed = json_parser.parse(fixtures_dir / "sample.json")
    email = "jane.doe@example.com"
    phone = "555-010-1234"
    email_block = next(b for b in parsed.blocks if b.id == "user.email")
    phone_block = next(b for b in parsed.blocks if b.id == "contacts[0].phone")

    dets = [
        _det("user.email", "email", email, 0),
        _det("contacts[0].phone", "phone", phone, 0),
    ]

    record = _make_record(tmp_path, fixtures_dir, "sample.json", FileType.JSON)

    out_path = _correct_json(
        record, {"user.email": [dets[0]], "contacts[0].phone": [dets[1]]}
    )

    import json as json_stdlib

    data = json_stdlib.loads(out_path.read_text(encoding="utf-8"))
    assert data["user"]["email"] == "X" * len(email)
    assert data["contacts"][0]["phone"] == "X" * len(phone)

    # Untouched fields survive verbatim.
    assert data["user"]["name"] == "Jane Doe"
    assert data["notes"] == "no pii here"
    assert data["active"] is True
    assert email_block.text == email  # sanity: fixture wasn't already masked
    assert phone_block.text == phone
